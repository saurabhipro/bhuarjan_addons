# -*- coding: utf-8 -*-

from odoo import models, fields, api, _
from odoo.exceptions import ValidationError
import base64
import io
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentTemplateMaster(models.Model):
    _name = 'bhu.document.template.master'
    _description = 'Document Template Master / दस्तावेज़ टेम्पलेट मास्टर'
    _inherit = ['mail.thread', 'mail.activity.mixin']
    _order = 'category, district_id, name'

    name = fields.Char(string='Template Name / टेम्पलेट नाम', required=True, tracking=True,
                      help='Name of the document template')
    
    category = fields.Selection([
        ('section4', 'Section 4 Notification / धारा 4 सूचना'),
        ('section8', 'Section 8 / धारा 8'),
        ('section11', 'Section 11 Preliminary Report / धारा 11 प्रारंभिक रिपोर्ट'),
        ('section15', 'Section 15 Objections / धारा 15 आपत्तियां'),
        ('section18', 'Section 18 R&R Scheme / धारा 18 पुनर्वास योजना'),
        ('section19', 'Section 19 Notification / धारा 19 सूचना'),
        ('section21', 'Section 21 Notification / धारा 21 सूचना'),
        ('section23', 'Section 23 Award / धारा 23 पुरस्कार'),
        ('sia', 'SIA Team / एसआईए टीम'),
        ('expert_group', 'Expert Group / विशेषज्ञ समूह'),
        ('form10', 'Form 10 / फॉर्म 10'),
        ('other', 'Other / अन्य'),
    ], string='Document Category / दस्तावेज़ श्रेणी', required=True, tracking=True,
       help='Category of the document template')
    
    district_id = fields.Many2one('bhu.district', string='District / जिला', 
                                  required=True, tracking=True,
                                  help='District for which this template is applicable')
    
    user_role_ids = fields.Many2many('res.groups', 'document_template_group_rel', 
                                     'template_id', 'group_id',
                                     string='User Roles / उपयोगकर्ता भूमिकाएं',
                                     tracking=True,
                                     help='User roles/groups that can access this template')
    
    template_content = fields.Html(string='Template Content / टेम्पलेट सामग्री', 
                                  required=True, tracking=False,
                                  help='Rich text template content in Hindi')
    
    active = fields.Boolean(string='Active / सक्रिय', default=True, tracking=True)
    
    # Computed fields for file names
    word_filename = fields.Char(string='Word Filename', compute='_compute_filenames', store=False)
    pdf_filename = fields.Char(string='PDF Filename', compute='_compute_filenames', store=False)
    
    @api.depends('name', 'category', 'district_id')
    def _compute_filenames(self):
        """Compute filenames for Word and PDF downloads"""
        for record in self:
            category_name = dict(record._fields['category'].selection).get(record.category, 'Template')
            district_name = record.district_id.name or 'District'
            safe_name = (record.name or 'Template').replace(' ', '_').replace('/', '_')
            record.word_filename = f"{category_name}_{district_name}_{safe_name}.docx"
            record.pdf_filename = f"{category_name}_{district_name}_{safe_name}.pdf"
    
    def action_download_word(self):
        """Download template as Word document"""
        self.ensure_one()
        
        if not DOCX_AVAILABLE:
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'title': _('Error'),
                    'message': _('python-docx library is not installed. Please install it using: pip install python-docx'),
                    'type': 'danger',
                    'sticky': False,
                }
            }
        
        try:
            # Create a new Word document
            doc = Document()
            
            # Set document title
            title = doc.add_heading(self.name or 'Document Template', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add category and district info
            category_name = dict(self._fields['category'].selection).get(self.category, 'N/A')
            doc.add_paragraph(f"Category / श्रेणी: {category_name}")
            doc.add_paragraph(f"District / जिला: {self.district_id.name or 'N/A'}")
            doc.add_paragraph('')  # Empty line
            
            # Convert HTML to plain text for Word (basic conversion)
            import re
            from html import unescape
            
            # Get plain text from HTML
            text_content = re.sub(r'<[^>]+>', '', self.template_content or '')
            text_content = unescape(text_content)
            
            # Add template content
            for paragraph in text_content.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    # Set font to support Hindi
                    for run in p.runs:
                        run.font.name = 'Mangal'
                        run.font.size = Pt(12)
            
            # Save to BytesIO
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            # Create attachment
            attachment = self.env['ir.attachment'].create({
                'name': self.word_filename,
                'type': 'binary',
                'datas': base64.b64encode(output.getvalue()),
                'res_model': 'bhu.document.template.master',
                'res_id': self.id,
                'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            })
            
            # Return download action
            return {
                'type': 'ir.actions.act_url',
                'url': f'/web/content/{attachment.id}?download=true',
                'target': 'self',
            }
        except Exception as e:
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'title': _('Error'),
                    'message': _('Error generating Word document: %s') % str(e),
                    'type': 'danger',
                    'sticky': False,
                }
            }
    
    def action_download_pdf(self):
        """Download template as PDF"""
        self.ensure_one()
        
        # Use Odoo's report engine to generate PDF
        report_action = self.env.ref('bhuarjan.action_report_document_template_pdf', raise_if_not_found=False)
        if report_action:
            return report_action.report_action(self)
        else:
            # Fallback: return a message
            return {
                'type': 'ir.actions.client',
                'tag': 'display_notification',
                'params': {
                    'title': _('PDF Report'),
                    'message': _('PDF report is being generated. Please configure the report action.'),
                    'type': 'warning',
                    'sticky': False,
                }
            }
    
    def action_edit_template(self):
        """Open template in edit mode"""
        self.ensure_one()
        return {
            'type': 'ir.actions.act_window',
            'name': _('Edit Document Template'),
            'res_model': 'bhu.document.template.master',
            'res_id': self.id,
            'view_mode': 'form',
            'target': 'current',
        }
    
    _sql_constraints = [
        ('unique_category_district', 'unique(category, district_id, name)', 
         'A template with the same category, district, and name already exists!')
    ]

